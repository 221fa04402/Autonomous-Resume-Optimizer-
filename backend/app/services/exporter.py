# backend/app/services/exporter.py
from docx import Document
import os

TMP = "tmp_uploads"
os.makedirs(TMP, exist_ok=True)

def make_docx(optimized_text, uid):
    # optimized_text can be a long string — just put it into sections
    doc = Document()
    doc.add_heading("Optimized Resume", level=1)
    for line in optimized_text.splitlines():
        if not line.strip():
            continue
        # simple heuristic: treat lines starting with '-' or '*' as bullets
        if line.strip().startswith(("-", "*")):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.strip().lstrip("-* ").strip())
        else:
            doc.add_paragraph(line.strip())
    filename = f"{uid}_optimized.docx"
    path = os.path.join(TMP, filename)
    doc.save(path)
    return path

def docx_to_pdf(docx_path):
    # Optional: convert via libreoffice (needs to be installed) — best run in backend container
    try:
        import subprocess, shutil
        if shutil.which("libreoffice") is None:
            return ""
        outdir = os.path.dirname(docx_path)
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", outdir
        ], check=True, timeout=30)
        base = os.path.splitext(os.path.basename(docx_path))[0]
        return os.path.join(outdir, base + ".pdf")
    except Exception as e:
        print("PDF conversion failed:", e)
        return ""
